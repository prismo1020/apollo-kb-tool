from __future__ import annotations

import argparse
import csv
import html
import json
import mimetypes
import re
import shutil
import threading
from collections import Counter
from datetime import datetime
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any
from urllib.parse import unquote, urlparse

from docx import Document


APP_DIR = Path(__file__).resolve().parent
KB_ROOT = APP_DIR.parent.resolve()
STATIC_DIR = APP_DIR / "static"
DATA_DIR = APP_DIR / "data"
BACKUP_ROOT = KB_ROOT / "_apollo_backups"
LOG_JSONL = DATA_DIR / "corrections.jsonl"
LOG_CSV = DATA_DIR / "corrections.csv"

APP_NAME = "Apollo Correction Console"

STOPWORDS = {
    "about",
    "after",
    "again",
    "all",
    "also",
    "and",
    "answer",
    "apollo",
    "are",
    "ask",
    "asked",
    "but",
    "can",
    "correct",
    "did",
    "does",
    "for",
    "from",
    "get",
    "had",
    "has",
    "have",
    "how",
    "into",
    "its",
    "just",
    "know",
    "like",
    "not",
    "our",
    "said",
    "say",
    "should",
    "that",
    "the",
    "their",
    "them",
    "then",
    "there",
    "this",
    "was",
    "what",
    "when",
    "where",
    "which",
    "with",
    "wrong",
    "you",
}

FORMATTING_RULES = [
    "*FORMATTING RULE FOR ALL APOLLO RESPONSES*",
    "Links MUST always be embedded: <URL|Display Text>. Never paste a raw URL.",
    "- *Bold* uses single asterisks. Never use **double asterisks**.",
    "- Do NOT use [text](url) markdown link format. Always use <url|text> Slack format.",
    "- Do NOT use # or ## headers in responses. Use *bold labels* instead.",
    "- Bullet lists use - or *. Numbered lists use 1. 2. 3.",
]

index_lock = threading.Lock()
write_lock = threading.Lock()
index_cache: dict[str, Any] = {"signature": None, "records": [], "built_at": None}


def now_stamp() -> str:
    return datetime.now().strftime("%Y%m%d-%H%M%S")


def now_iso() -> str:
    return datetime.now().isoformat(timespec="seconds")


def tokenize(text: str) -> list[str]:
    words = re.findall(r"[a-zA-Z0-9]{2,}", text.lower())
    return [word for word in words if word not in STOPWORDS]


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value).strip())


def safe_multiline(value: Any) -> str:
    if value is None:
        return ""
    lines = [re.sub(r"\s+", " ", line).strip() for line in str(value).splitlines()]
    return "\n".join(line for line in lines if line)


def slugify(value: str, fallback: str = "NEW_CORRECTION") -> str:
    tokens = re.findall(r"[A-Za-z0-9]+", value.upper())
    slug = "_".join(tokens[:8])
    return slug[:80] or fallback


def is_heading(text: str) -> bool:
    stripped = clean_text(text)
    if len(stripped) < 3 or len(stripped) > 140:
        return False
    if not (stripped.startswith("*") and stripped.endswith("*")):
        return False
    if stripped.startswith("-") or stripped.startswith("•"):
        return False
    inner = stripped.strip("*").strip()
    return bool(inner) and not inner.endswith(":")


def is_separator(text: str) -> bool:
    return clean_text(text) in {"---", "--", "___"}


def normalize_brand_terms(text: str) -> str:
    replacements = {
        r"\btact\s*suit\s*x40\b": "TactSuit X40",
        r"\btactsuit\s*x40\b": "TactSuit X40",
        r"\btact\s*suit\s*pro\b": "TactSuit Pro",
        r"\btactsuit\s*pro\b": "TactSuit Pro",
        r"\bvest workflow\b": "Vest Workflow",
        r"\bt1\b": "T1",
    }
    result = text
    for pattern, replacement in replacements.items():
        result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)
    return result


def category_from_filename(path: Path) -> str:
    stem = path.stem.strip("_")
    parts = stem.split("__")
    if len(parts) >= 2:
        left = parts[0]
        maybe = left.split("_")[-1]
        if maybe and not maybe.isdigit():
            return maybe.upper()
    for category in [
        "TECH",
        "OPERATIONS",
        "POLICY",
        "EXPERIENCES",
        "SYSTEMS",
        "ADMIN",
        "STAFF",
        "GUEST",
        "VISUAL",
        "OPS_FLOW",
        "CORRECTION",
    ]:
        if category in stem.upper():
            return category
    return "GENERAL"


def is_safe_kb_path(path: Path) -> bool:
    resolved = path.resolve()
    if not resolved.is_relative_to(KB_ROOT):
        return False
    if resolved.is_relative_to(APP_DIR):
        return False
    if resolved.is_relative_to(BACKUP_ROOT):
        return False
    parts = resolved.relative_to(KB_ROOT).parts
    if "brand-preview" in parts:
        return False
    if "Master List" in parts:
        return False
    return True


def collect_docx_files() -> list[Path]:
    files: list[Path] = []
    for path in KB_ROOT.rglob("*.docx"):
        if path.name.startswith("~$"):
            continue
        if is_safe_kb_path(path):
            files.append(path)
    return sorted(files, key=lambda p: str(p.relative_to(KB_ROOT)).lower())


def current_signature(files: list[Path]) -> tuple[tuple[str, int, int], ...]:
    return tuple(
        (str(path.relative_to(KB_ROOT)), path.stat().st_mtime_ns, path.stat().st_size)
        for path in files
    )


def document_text(path: Path) -> tuple[str, list[str]]:
    doc = Document(str(path))
    chunks: list[str] = []
    paragraphs: list[str] = []

    for paragraph in doc.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            paragraphs.append(text)
            chunks.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(clean_text(cell.text) for cell in row.cells if clean_text(cell.text))
            if row_text:
                paragraphs.append(row_text)
                chunks.append(row_text)

    return "\n".join(chunks), paragraphs


def parse_sections(path: Path) -> list[dict[str, Any]]:
    doc = Document(str(path))
    headings = [
        index
        for index, paragraph in enumerate(doc.paragraphs)
        if is_heading(paragraph.text)
    ]
    sections: list[dict[str, Any]] = []
    for pos, start in enumerate(headings):
        next_start = headings[pos + 1] if pos + 1 < len(headings) else len(doc.paragraphs)
        end = next_start

        while end > start + 1 and not clean_text(doc.paragraphs[end - 1].text):
            end -= 1
        if end > start + 1 and is_separator(doc.paragraphs[end - 1].text):
            end -= 1
        while end > start + 1 and not clean_text(doc.paragraphs[end - 1].text):
            end -= 1

        lines = [
            clean_text(doc.paragraphs[index].text)
            for index in range(start, end)
            if clean_text(doc.paragraphs[index].text)
        ]
        heading = clean_text(doc.paragraphs[start].text)
        body = "\n".join(lines[1:])
        text = "\n".join(lines)
        tokens = tokenize(f"{heading}\n{body}")
        sections.append(
            {
                "heading": heading,
                "start": start,
                "end": end,
                "text": text,
                "body": body,
                "tokens": Counter(tokens),
                "token_set": set(tokens),
                "heading_tokens": set(tokenize(heading)),
            }
        )
    return sections


def load_index(force: bool = False) -> list[dict[str, Any]]:
    files = collect_docx_files()
    signature = current_signature(files)

    with index_lock:
        if not force and index_cache["signature"] == signature:
            return index_cache["records"]

        records: list[dict[str, Any]] = []
        for path in files:
            rel = path.relative_to(KB_ROOT)
            try:
                text, paragraphs = document_text(path)
                sections = parse_sections(path)
            except Exception as exc:
                text = ""
                paragraphs = [f"Unable to read document: {exc}"]
                sections = []

            name_text = path.stem.replace("_", " ")
            tokens = tokenize(f"{name_text}\n{text}")
            records.append(
                {
                    "file": rel.as_posix(),
                    "name": path.name,
                    "category": category_from_filename(path),
                    "tokens": Counter(tokens),
                    "token_set": set(tokens),
                    "name_tokens": set(tokenize(name_text)),
                    "paragraphs": paragraphs[:300],
                    "sections": sections,
                    "char_count": len(text),
                    "modified": datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds"),
                }
            )

        index_cache["signature"] = signature
        index_cache["records"] = records
        index_cache["built_at"] = now_iso()
        return records


def best_snippet(record: dict[str, Any], query_tokens: set[str]) -> str:
    best = ""
    best_score = 0
    for paragraph in record["paragraphs"]:
        ptokens = set(tokenize(paragraph))
        score = len(ptokens & query_tokens)
        if score > best_score:
            best = paragraph
            best_score = score
    if not best and record["paragraphs"]:
        best = record["paragraphs"][0]
    if len(best) > 320:
        return best[:317].rstrip() + "..."
    return best


def best_section(record: dict[str, Any], query_tokens: set[str]) -> dict[str, Any] | None:
    best: dict[str, Any] | None = None
    best_score = 0.0
    for section in record.get("sections", []):
        overlap = query_tokens & section["token_set"]
        if not overlap:
            continue
        score = 0.0
        for token in overlap:
            score += 1.0 + min(section["tokens"][token], 4) * 0.2
            if token in section["heading_tokens"]:
                score += 2.8
        if score > best_score:
            best_score = score
            best = section

    if best is None and record.get("sections"):
        best = record["sections"][0]
        best_score = 0.0

    if best is None:
        return None

    return {
        "heading": best["heading"],
        "text": best["text"],
        "score": round(best_score, 2),
        "confidence": confidence_label(best_score),
    }


def approved_guidance(payload: dict[str, Any]) -> str:
    guidance = safe_multiline(payload.get("correct_answer"))
    if guidance:
        return normalize_brand_terms(guidance)
    return (
        "Apollo should not provide a definitive answer from the current KB. "
        "Escalate for verification or ask for the missing source before answering."
    )


def looks_like_full_section(text: str) -> bool:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return bool(lines) and is_heading(lines[0]) and len(lines) > 1


def build_led_bullet(payload: dict[str, Any]) -> str | None:
    combined = " ".join(
        [
            safe_multiline(payload.get("question")),
            safe_multiline(payload.get("wrong_answer")),
            safe_multiline(payload.get("correct_answer")),
            safe_multiline(payload.get("notes")),
        ]
    ).lower()
    has_red_blue = "red and blue" in combined or "red/blue" in combined
    has_vest = any(term in combined for term in ["vest", "tactsuit", "tact suit", "x40"])
    if not (has_red_blue and has_vest):
        return None
    if "low battery" not in combined and "battery" not in combined:
        return None
    return (
        "- Flashing red and blue = low battery on the TactSuit X40. Charge the vest. "
        "If a TactSuit Pro flashes red and blue, this should never happen -- immediately submit a Vest Workflow."
    )


def normalized_led_key(line: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", line.lower()).strip()


def merge_led_guidance(section_text: str, payload: dict[str, Any]) -> str | None:
    led_bullet = build_led_bullet(payload)
    if not led_bullet:
        return None

    lines = [line.rstrip() for line in section_text.splitlines() if line.strip()]
    if not lines:
        return None

    replaced = False
    insert_at: int | None = None
    output: list[str] = []
    for line in lines:
        key = normalized_led_key(line)
        key_words = set(key.split())
        if "red" in key_words and "blue" in key_words:
            if not replaced:
                output.append(led_bullet)
                replaced = True
            continue
        output.append(line)
        if insert_at is None and "flashing blue" in key:
            insert_at = len(output)

    if not replaced:
        if insert_at is None:
            insert_at = 1 if len(output) > 1 else len(output)
        output.insert(insert_at, led_bullet)

    return "\n".join(output)


def build_section_replacement(section_text: str, payload: dict[str, Any]) -> str:
    guidance = approved_guidance(payload)
    if looks_like_full_section(guidance):
        return guidance

    led_merge = merge_led_guidance(section_text, payload)
    if led_merge:
        return led_merge

    lines = [line.rstrip() for line in section_text.splitlines() if line.strip()]
    if not lines:
        return f"*Approved Guidance*\n- {guidance}"

    bullet = guidance if guidance.startswith("- ") else f"- {guidance}"
    if bullet not in lines:
        lines.append(bullet)
    return "\n".join(lines)


def build_new_file_lines(payload: dict[str, Any]) -> list[str]:
    guidance = approved_guidance(payload)
    if looks_like_full_section(guidance):
        return [line for line in guidance.splitlines() if line.strip()]
    return build_correction_lines({**payload, "correct_answer": guidance})


def find_matches(payload: dict[str, Any], limit: int = 8) -> list[dict[str, Any]]:
    records = load_index()
    query = "\n".join(
        [
            safe_multiline(payload.get("question")),
            safe_multiline(payload.get("wrong_answer")),
            safe_multiline(payload.get("correct_answer")),
            safe_multiline(payload.get("notes")),
            clean_text(payload.get("category")),
        ]
    )
    query_counts = Counter(tokenize(query))
    query_tokens = set(query_counts)
    if not query_tokens:
        return []

    matches: list[dict[str, Any]] = []
    requested_category = clean_text(payload.get("category")).upper()
    for record in records:
        overlap = query_tokens & record["token_set"]
        if not overlap:
            continue
        weighted = 0.0
        for token in overlap:
            base = 1.0 + min(record["tokens"][token], 6) * 0.12
            if token in record["name_tokens"]:
                base += 2.2
            weighted += base
        if requested_category and requested_category == record["category"]:
            weighted += 2.5
        section = best_section(record, query_tokens)
        if section:
            weighted += min(section["score"], 8)
        score = round(weighted, 2)
        section_text = section["text"] if section else ""
        matches.append(
            {
                "file": record["file"],
                "name": record["name"],
                "category": record["category"],
                "score": score,
                "confidence": confidence_label(score),
                "snippet": best_snippet(record, query_tokens),
                "section_heading": section["heading"] if section else "",
                "section_confidence": section["confidence"] if section else "Low",
                "section_score": section["score"] if section else 0,
                "section_text": section_text,
                "proposed_text": build_section_replacement(section_text, payload) if section_text else draft_preview(payload),
                "modified": record["modified"],
            }
        )

    matches.sort(key=lambda item: item["score"], reverse=True)
    return matches[:limit]


def confidence_label(score: float) -> str:
    if score >= 20:
        return "High"
    if score >= 10:
        return "Medium"
    return "Low"


def proposed_action(matches: list[dict[str, Any]]) -> str:
    if not matches:
        return "new_file"
    if matches[0]["score"] < 8:
        return "new_file"
    return "update_existing"


def build_correction_lines(payload: dict[str, Any]) -> list[str]:
    question = safe_multiline(payload.get("question"))
    wrong = safe_multiline(payload.get("wrong_answer"))
    correct = safe_multiline(payload.get("correct_answer"))
    notes = safe_multiline(payload.get("notes"))
    submitter = clean_text(payload.get("submitter")) or "Unspecified reviewer"

    if not correct:
        correct = (
            "Apollo should not provide a definitive answer from the current KB. "
            "Escalate for verification or ask for the missing source before answering."
        )

    lines = [
        "*Approved Guidance*",
        f"*Submitted by:* {submitter}",
    ]
    if question:
        lines.extend(["*Use this guidance when the user asks:*", question])
    if wrong:
        lines.extend(["*Do not answer this way:*", wrong])
    lines.extend(["*Apollo should say:*", normalize_brand_terms(correct)])
    if notes:
        lines.extend(["*Reviewer notes:*", notes])
    lines.append("*Operational rule:* Prefer this guidance over older conflicting guidance.")
    return lines


def draft_preview(payload: dict[str, Any]) -> str:
    return "\n".join(build_correction_lines(payload))


def next_file_number() -> int:
    highest = 0
    for path in collect_docx_files():
        match = re.match(r"_?(\d+)[_-]", path.name)
        if match:
            highest = max(highest, int(match.group(1)))
    return highest + 1


def build_new_filename(payload: dict[str, Any]) -> str:
    category = slugify(clean_text(payload.get("category")) or "CORRECTION", "CORRECTION")
    topic_source = clean_text(payload.get("new_topic")) or clean_text(payload.get("question")) or "NEW_CORRECTION"
    topic = slugify(topic_source, "NEW_CORRECTION")
    number = next_file_number()
    return f"_{number:02d}_{category}__{topic}.docx"


def add_lines_to_doc(doc: Document, lines: list[str]) -> None:
    for line in lines:
        if "\n" in line:
            for part in line.splitlines():
                doc.add_paragraph(part)
        else:
            doc.add_paragraph(line)


def replacement_lines(text: str, fallback_heading: str = "") -> list[str]:
    lines = [clean_text(line) for line in str(text or "").splitlines() if clean_text(line)]
    if fallback_heading and (not lines or not is_heading(lines[0])):
        lines.insert(0, fallback_heading)
    return lines


def delete_paragraph(paragraph: Any) -> None:
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)
    paragraph._p = paragraph._element = None


def section_bounds_in_doc(doc: Document, heading: str) -> tuple[int, int]:
    clean_heading = clean_text(heading)
    if not clean_heading:
        raise ValueError("No section heading selected.")

    heading_indices = [
        index
        for index, paragraph in enumerate(doc.paragraphs)
        if is_heading(paragraph.text)
    ]
    for position, start in enumerate(heading_indices):
        if clean_text(doc.paragraphs[start].text) != clean_heading:
            continue
        next_start = heading_indices[position + 1] if position + 1 < len(heading_indices) else len(doc.paragraphs)
        end = next_start
        while end > start + 1 and not clean_text(doc.paragraphs[end - 1].text):
            end -= 1
        if end > start + 1 and is_separator(doc.paragraphs[end - 1].text):
            end -= 1
        while end > start + 1 and not clean_text(doc.paragraphs[end - 1].text):
            end -= 1
        return start, end

    raise ValueError(f"Could not find section heading: {clean_heading}")


def section_text_from_doc(doc: Document, start: int, end: int) -> str:
    return "\n".join(
        clean_text(doc.paragraphs[index].text)
        for index in range(start, end)
        if clean_text(doc.paragraphs[index].text)
    )


def create_new_docx(target: Path, payload: dict[str, Any], content_lines: list[str]) -> None:
    doc = Document()
    title = target.stem.strip("_")
    doc.add_paragraph(f"*{title}*")
    add_lines_to_doc(doc, FORMATTING_RULES)
    doc.add_paragraph("---")
    doc.add_paragraph("*Purpose*")
    purpose = clean_text(payload.get("new_purpose")) or (
        "This file captures approved Apollo guidance created from a reviewed correction."
    )
    doc.add_paragraph(purpose)
    doc.add_paragraph("*When to Use This File*")
    question = clean_text(payload.get("question"))
    if question:
        doc.add_paragraph(f"Use this file when the user asks about: {question}")
    else:
        doc.add_paragraph("Use this file when the user's message matches the approved guidance below.")
    add_lines_to_doc(doc, content_lines)
    doc.save(str(target))


def replace_existing_section_docx(
    target: Path,
    section_heading: str,
    replacement_text: str,
) -> tuple[Path, str, str]:
    backup_dir = BACKUP_ROOT / datetime.now().strftime("%Y%m%d")
    backup_dir.mkdir(parents=True, exist_ok=True)
    backup = backup_dir / f"{target.stem}__backup_{now_stamp()}{target.suffix}"
    shutil.copy2(target, backup)

    doc = Document(str(target))
    start, end = section_bounds_in_doc(doc, section_heading)
    old_text = section_text_from_doc(doc, start, end)
    lines = replacement_lines(replacement_text, section_heading)
    if len(lines) < 2:
        raise ValueError("Replacement section needs a heading and at least one guidance line.")

    reference = doc.paragraphs[start]
    for line in lines:
        reference.insert_paragraph_before(line)

    old_paragraphs = list(doc.paragraphs[start + len(lines) : end + len(lines)])
    for paragraph in old_paragraphs:
        delete_paragraph(paragraph)

    doc.save(str(target))
    return backup, old_text, "\n".join(lines)


def write_log(row: dict[str, Any]) -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    with LOG_JSONL.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(row, ensure_ascii=True) + "\n")

    write_header = not LOG_CSV.exists()
    with LOG_CSV.open("a", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(
            handle,
            fieldnames=[
                "id",
                "timestamp",
                "mode",
                "target_file",
                "backup_file",
                "section_heading",
                "submitter",
                "question",
                "status",
            ],
        )
        if write_header:
            writer.writeheader()
        writer.writerow(
            {
                "id": row["id"],
                "timestamp": row["timestamp"],
                "mode": row["mode"],
                "target_file": row["target_file"],
                "backup_file": row.get("backup_file", ""),
                "section_heading": row.get("section_heading", ""),
                "submitter": row.get("submitter", ""),
                "question": row.get("question", ""),
                "status": row["status"],
            }
        )


def recent_log(limit: int = 10) -> list[dict[str, Any]]:
    if not LOG_JSONL.exists():
        return []
    lines = LOG_JSONL.read_text(encoding="utf-8").splitlines()[-limit:]
    items = []
    for line in lines:
        try:
            items.append(json.loads(line))
        except json.JSONDecodeError:
            continue
    return list(reversed(items))


def approve_update(payload: dict[str, Any]) -> dict[str, Any]:
    mode = clean_text(payload.get("mode")) or "existing"
    submitter = clean_text(payload.get("submitter")) or "Unspecified reviewer"
    old_section = ""
    new_section = ""

    with write_lock:
        if mode == "new":
            filename = clean_text(payload.get("new_filename")) or build_new_filename(payload)
            filename = filename if filename.lower().endswith(".docx") else f"{filename}.docx"
            target = (KB_ROOT / Path(filename).name).resolve()
            if not is_safe_kb_path(target):
                raise ValueError("Unsafe new file path.")
            if target.exists():
                raise FileExistsError(f"{target.name} already exists.")
            lines = replacement_lines(payload.get("replacement_text")) or build_new_file_lines(payload)
            create_new_docx(target, payload, lines)
            new_section = "\n".join(lines)
            backup = None
        else:
            rel = clean_text(payload.get("target_file"))
            if not rel:
                raise ValueError("No target file selected.")
            target = (KB_ROOT / unquote(rel)).resolve()
            if not is_safe_kb_path(target) or not target.exists() or target.suffix.lower() != ".docx":
                raise ValueError("Selected target file is not a valid KB document.")
            section_heading = clean_text(payload.get("section_heading"))
            replacement_text = safe_multiline(payload.get("replacement_text"))
            if not replacement_text:
                raise ValueError("No replacement section text was submitted.")
            backup, old_section, new_section = replace_existing_section_docx(
                target,
                section_heading,
                replacement_text,
            )

        row = {
            "id": now_stamp(),
            "timestamp": now_iso(),
            "mode": "new_file" if mode == "new" else "replace_section",
            "target_file": target.relative_to(KB_ROOT).as_posix(),
            "backup_file": backup.relative_to(KB_ROOT).as_posix() if backup else "",
            "section_heading": clean_text(payload.get("section_heading")),
            "submitter": submitter,
            "question": clean_text(payload.get("question")),
            "wrong_answer": safe_multiline(payload.get("wrong_answer")),
            "correct_answer": safe_multiline(payload.get("correct_answer")),
            "notes": safe_multiline(payload.get("notes")),
            "old_section": old_section,
            "new_section": new_section,
            "status": "approved_written",
        }
        write_log(row)
        load_index(force=True)
        return row


class AppHandler(BaseHTTPRequestHandler):
    server_version = "ApolloCorrectionServer/0.1"

    def log_message(self, format: str, *args: Any) -> None:
        return

    def _send_json(self, payload: Any, status: int = 200) -> None:
        raw = json.dumps(payload, ensure_ascii=True).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(raw)))
        self.send_header("Cache-Control", "no-store")
        self.end_headers()
        self.wfile.write(raw)

    def _read_json(self) -> dict[str, Any]:
        length = int(self.headers.get("Content-Length", "0"))
        if length <= 0:
            return {}
        raw = self.rfile.read(length).decode("utf-8")
        return json.loads(raw)

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        if parsed.path == "/api/status":
            records = index_cache["records"]
            file_count = len(records) if records else len(collect_docx_files())
            self._send_json(
                {
                    "app": APP_NAME,
                    "kb_root": str(KB_ROOT),
                    "file_count": file_count,
                    "built_at": index_cache["built_at"],
                    "index_state": "ready" if records else "indexing",
                    "recent": recent_log(),
                }
            )
            return

        if parsed.path == "/api/files":
            records = load_index()
            self._send_json(
                [
                    {
                        "file": record["file"],
                        "name": record["name"],
                        "category": record["category"],
                        "modified": record["modified"],
                        "char_count": record["char_count"],
                    }
                    for record in records
                ]
            )
            return

        self.serve_static(parsed.path)

    def do_POST(self) -> None:
        parsed = urlparse(self.path)
        try:
            payload = self._read_json()
            if parsed.path == "/api/analyze":
                matches = find_matches(payload)
                self._send_json(
                    {
                        "matches": matches,
                        "suggested_action": proposed_action(matches),
                        "draft": draft_preview(payload),
                    }
                )
                return

            if parsed.path == "/api/approve":
                result = approve_update(payload)
                self._send_json({"ok": True, "result": result})
                return

            self._send_json({"error": "Unknown endpoint."}, status=404)
        except Exception as exc:
            self._send_json({"error": str(exc)}, status=400)

    def serve_static(self, path: str) -> None:
        if path in ("", "/"):
            path = "/index.html"
        rel = Path(unquote(path.lstrip("/")))
        target = (STATIC_DIR / rel).resolve()
        if not target.is_relative_to(STATIC_DIR) or not target.exists() or not target.is_file():
            self.send_response(404)
            self.send_header("Content-Type", "text/plain; charset=utf-8")
            self.end_headers()
            self.wfile.write(b"Not found")
            return

        raw = target.read_bytes()
        content_type = mimetypes.guess_type(str(target))[0] or "application/octet-stream"
        if target.suffix.lower() == ".html":
            content_type = "text/html; charset=utf-8"
        elif target.suffix.lower() == ".css":
            content_type = "text/css; charset=utf-8"
        elif target.suffix.lower() == ".js":
            content_type = "text/javascript; charset=utf-8"

        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(raw)))
        self.end_headers()
        self.wfile.write(raw)


def main() -> None:
    parser = argparse.ArgumentParser(description=APP_NAME)
    parser.add_argument("--host", default="127.0.0.1")
    parser.add_argument("--port", type=int, default=8787)
    args = parser.parse_args()

    DATA_DIR.mkdir(parents=True, exist_ok=True)

    server = ThreadingHTTPServer((args.host, args.port), AppHandler)
    threading.Thread(target=load_index, kwargs={"force": True}, daemon=True).start()
    print(f"{APP_NAME} running at http://{args.host}:{args.port}", flush=True)
    print(f"Knowledge base: {html.escape(str(KB_ROOT))}", flush=True)
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        pass
    finally:
        server.server_close()


if __name__ == "__main__":
    main()
